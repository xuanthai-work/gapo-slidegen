from datetime import timedelta
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.dialects import postgresql

from app.auth import get_current_user
from app.ingestion import SourceDocument
from app.main import app
from app.models import SourceRecord, User
from app.sources.dependencies import get_source_service
from app.sources.service import SourceService, build_owned_source_query, build_owned_sources_query
from app.storage import LocalObjectStorage


class FakeSession:
    def __init__(self, fail_flush: bool = False) -> None:
        self.added: list[object] = []
        self.fail_flush = fail_flush

    def add(self, value: object) -> None:
        self.added.append(value)

    def flush(self) -> None:
        if self.fail_flush:
            raise RuntimeError("database unavailable")


def _user() -> User:
    return User(
        id=uuid4(),
        email="owner@example.com",
        normalized_email="owner@example.com",
        password_hash="not-used",
        is_active=True,
    )


def _docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_heading("Source title", level=1)
    document.add_paragraph("Content owned by one internal user.")
    document.save(stream)
    return stream.getvalue()


def test_text_source_is_bound_to_authenticated_owner(tmp_path: Path) -> None:
    session = FakeSession()
    service = SourceService(session, LocalObjectStorage(tmp_path), timedelta(hours=24))  # type: ignore[arg-type]
    user = _user()
    document = SourceDocument(
        kind="prompt",
        title="Internal plan",
        text="Build a concise deck",
        sections=[{"index": 0, "title": "Internal plan", "text": "Build a concise deck"}],
    )

    record = service.create_text(user, document)

    assert record.owner_id == user.id
    assert record.title == "Internal plan"
    assert record.storage_key is None
    assert record.delete_after is not None


def test_file_source_is_stored_below_owner_prefix(tmp_path: Path) -> None:
    session = FakeSession()
    storage = LocalObjectStorage(tmp_path)
    service = SourceService(session, storage, timedelta(hours=24))  # type: ignore[arg-type]
    user = _user()
    data = _docx_bytes()

    record = service.create_file(
        user,
        filename="../../brief.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=data,
    )

    assert record.storage_key is not None
    assert record.storage_key.startswith(f"users/{user.id}/sources/")
    assert storage.get(record.storage_key) == data
    assert "Content owned" in record.extracted_text


def test_file_is_removed_when_database_flush_fails(tmp_path: Path) -> None:
    service = SourceService(
        FakeSession(fail_flush=True),  # type: ignore[arg-type]
        LocalObjectStorage(tmp_path),
        timedelta(hours=24),
    )
    with pytest.raises(RuntimeError, match="database unavailable"):
        service.create_file(
            _user(),
            filename="brief.docx",
            content_type=None,
            data=_docx_bytes(),
        )
    assert not any(path.is_file() for path in tmp_path.rglob("*"))


def test_owned_source_query_filters_both_source_and_owner() -> None:
    source_id = uuid4()
    owner_id = uuid4()
    sql = str(
        build_owned_source_query(source_id, owner_id).compile(
            dialect=postgresql.dialect(), compile_kwargs={"literal_binds": True}
        )
    )
    assert str(source_id) in sql
    assert str(owner_id) in sql
    assert "source_records.owner_id" in sql


def test_source_list_query_is_owned_and_bounded() -> None:
    owner_id = uuid4()
    sql = str(
        build_owned_sources_query(owner_id, limit=50).compile(
            dialect=postgresql.dialect(), compile_kwargs={"literal_binds": True}
        )
    )
    assert str(owner_id) in sql
    assert "source_records.owner_id" in sql
    assert "LIMIT 50" in sql


def test_source_http_response_does_not_expose_storage_or_owner(tmp_path: Path) -> None:
    session = FakeSession()
    user = _user()
    service = SourceService(session, LocalObjectStorage(tmp_path), timedelta(hours=24))  # type: ignore[arg-type]
    app.dependency_overrides[get_current_user] = lambda: user
    app.dependency_overrides[get_source_service] = lambda: service
    client = TestClient(app)
    try:
        response = client.post(
            "/v1/sources/text",
            json={"kind": "manuscript", "title": "Private brief", "text": "Internal content"},
        )
        assert response.status_code == 201
        body = response.json()
        assert body["title"] == "Private brief"
        assert "owner_id" not in body
        assert "storage_key" not in body
    finally:
        app.dependency_overrides.clear()
